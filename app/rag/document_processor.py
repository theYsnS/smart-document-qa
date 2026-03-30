"""Document loading, preprocessing, and chunking."""

from __future__ import annotations

import hashlib
import logging
import re
from pathlib import Path

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from app.config import settings

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}


def generate_doc_id(filepath: Path) -> str:
    """Generate a deterministic document ID from file content hash."""
    content = filepath.read_bytes()
    return hashlib.sha256(content).hexdigest()[:16]


def _load_pdf(filepath: Path) -> list[Document]:
    """Load a PDF file and return one Document per page."""
    from pypdf import PdfReader

    reader = PdfReader(str(filepath))
    documents: list[Document] = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": filepath.name,
                        "page": page_num,
                        "total_pages": len(reader.pages),
                    },
                )
            )
    return documents


def _load_docx(filepath: Path) -> list[Document]:
    """Load a DOCX file and return a single Document."""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(filepath))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n\n".join(paragraphs)
    return [
        Document(
            page_content=full_text,
            metadata={"source": filepath.name},
        )
    ]


def _load_txt(filepath: Path) -> list[Document]:
    """Load a plain text file."""
    text = filepath.read_text(encoding="utf-8", errors="replace")
    return [
        Document(
            page_content=text,
            metadata={"source": filepath.name},
        )
    ]


_LOADERS = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".txt": _load_txt,
}


def clean_text(text: str) -> str:
    """Normalize whitespace and remove control characters."""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def load_document(filepath: Path) -> list[Document]:
    """Load a document from disk based on its file extension.

    Args:
        filepath: Path to the document file.

    Returns:
        List of LangChain Document objects.

    Raises:
        ValueError: If the file extension is not supported.
        FileNotFoundError: If the file does not exist.
    """
    if not filepath.exists():
        raise FileNotFoundError(f"File not found: {filepath}")

    ext = filepath.suffix.lower()
    loader = _LOADERS.get(ext)
    if loader is None:
        raise ValueError(
            f"Unsupported file type: {ext}. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    logger.info("Loading document: %s", filepath.name)
    documents = loader(filepath)

    # Clean text content
    for doc in documents:
        doc.page_content = clean_text(doc.page_content)

    logger.info("Loaded %d page(s) from %s", len(documents), filepath.name)
    return documents


def chunk_documents(
    documents: list[Document],
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[Document]:
    """Split documents into smaller chunks for embedding.

    Uses RecursiveCharacterTextSplitter with separators optimized for
    natural text boundaries (paragraphs, sentences, words).

    Args:
        documents: List of documents to chunk.
        chunk_size: Maximum chunk size in characters.
        chunk_overlap: Overlap between consecutive chunks.

    Returns:
        List of chunked Document objects with preserved metadata.
    """
    size = chunk_size or settings.chunk_size
    overlap = chunk_overlap or settings.chunk_overlap

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=size,
        chunk_overlap=overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", ", ", " ", ""],
        keep_separator=True,
    )

    chunks = splitter.split_documents(documents)

    # Add chunk index to metadata
    for idx, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = idx

    logger.info(
        "Split %d document(s) into %d chunks (size=%d, overlap=%d)",
        len(documents),
        len(chunks),
        size,
        overlap,
    )
    return chunks


def process_document(
    filepath: Path,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> tuple[str, list[Document]]:
    """Full pipeline: load, clean, chunk a document.

    Args:
        filepath: Path to the document file.
        chunk_size: Optional override for chunk size.
        chunk_overlap: Optional override for chunk overlap.

    Returns:
        Tuple of (doc_id, list of chunked documents).
    """
    doc_id = generate_doc_id(filepath)
    raw_docs = load_document(filepath)
    chunks = chunk_documents(raw_docs, chunk_size, chunk_overlap)

    # Inject doc_id into all chunk metadata
    for chunk in chunks:
        chunk.metadata["doc_id"] = doc_id

    return doc_id, chunks
